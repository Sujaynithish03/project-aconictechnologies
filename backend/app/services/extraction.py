"""File validation and text extraction for PDF, DOCX, and TXT uploads."""

import io
import logging
import re
from pathlib import Path

import pypdf
from docx import Document as DocxDocument

from app.core.config import settings
from app.core.exceptions import (
    EmptyFileError,
    FileTooLargeError,
    TextExtractionError,
    UnsupportedFileTypeError,
)

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {"pdf", "docx", "txt"}

# Accepted content types per extension. Browsers are inconsistent here, so the
# generic octet-stream is tolerated and the magic-byte check does the real work.
ALLOWED_CONTENT_TYPES = {
    "pdf": {"application/pdf", "application/octet-stream"},
    "docx": {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/zip",
        "application/octet-stream",
    },
    "txt": {"text/plain", "text/markdown", "application/octet-stream", ""},
}

PDF_MAGIC = b"%PDF-"
DOCX_MAGIC = b"PK\x03\x04"  # DOCX is a zip archive.

_WHITESPACE_RUN = re.compile(r"[ \t]{2,}")
_BLANK_LINES = re.compile(r"\n{3,}")


def resolve_extension(filename: str) -> str:
    extension = Path(filename).suffix.lower().lstrip(".")
    if extension not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '.{extension or filename}'. "
            "Upload a PDF, DOCX, or TXT file."
        )
    return extension


def validate_upload(filename: str, content_type: str | None, data: bytes) -> str:
    """Validate an upload on extension, declared MIME type, and magic bytes.

    Returns the normalised extension. Raises the matching ``AppError`` subclass
    so routes map cleanly onto 413/415/400.
    """
    extension = resolve_extension(filename)

    if not data:
        raise EmptyFileError()
    if len(data) > settings.max_upload_bytes:
        raise FileTooLargeError(
            f"File is {len(data) / 1_048_576:.1f} MB; the limit is "
            f"{settings.max_upload_mb} MB."
        )

    declared = (content_type or "").split(";")[0].strip().lower()
    if declared and declared not in ALLOWED_CONTENT_TYPES[extension]:
        raise UnsupportedFileTypeError(
            f"File content type '{declared}' does not match a .{extension} file."
        )

    # Magic bytes catch a renamed executable masquerading as a document.
    if extension == "pdf" and not data.startswith(PDF_MAGIC):
        raise UnsupportedFileTypeError("File does not appear to be a valid PDF.")
    if extension == "docx" and not data.startswith(DOCX_MAGIC):
        raise UnsupportedFileTypeError(
            "File does not appear to be a valid DOCX. Legacy .doc is not supported."
        )
    if extension == "txt":
        try:
            data.decode("utf-8")
        except UnicodeDecodeError:
            try:
                data.decode("latin-1")
            except UnicodeDecodeError as error:
                raise UnsupportedFileTypeError(
                    "Text file is not valid UTF-8 or Latin-1."
                ) from error

    return extension


def extract_text(data: bytes, extension: str) -> str:
    """Extract plain text from validated file bytes."""
    extractors = {
        "pdf": _extract_pdf,
        "docx": _extract_docx,
        "txt": _extract_txt,
    }
    extractor = extractors.get(extension)
    if extractor is None:
        raise UnsupportedFileTypeError()

    text = _normalise(extractor(data))
    if not text.strip():
        raise TextExtractionError(
            "No readable text found. Scanned or image-only documents are not "
            "supported because OCR is not enabled."
        )
    return text


def _extract_pdf(data: bytes) -> str:
    try:
        reader = pypdf.PdfReader(io.BytesIO(data))
    except Exception as error:
        raise TextExtractionError(f"Could not read the PDF: {error}") from error

    if reader.is_encrypted:
        # An empty user password is common; try it before giving up.
        try:
            reader.decrypt("")
        except Exception as error:
            raise TextExtractionError(
                "This PDF is password protected."
            ) from error

    pages: list[str] = []
    for number, page in enumerate(reader.pages, start=1):
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            # One malformed page shouldn't discard the whole document.
            logger.warning("Skipping unreadable PDF page %s", number)
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    try:
        document = DocxDocument(io.BytesIO(data))
    except Exception as error:
        raise TextExtractionError(f"Could not read the DOCX: {error}") from error

    parts = [p.text for p in document.paragraphs if p.text.strip()]
    # Tables often hold the substantive content (dates, policies, pricing).
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _extract_txt(data: bytes) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("latin-1", errors="replace")


def _normalise(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "")
    text = _WHITESPACE_RUN.sub(" ", text)
    return _BLANK_LINES.sub("\n\n", text).strip()
