"""Unit tests for the pure text-processing services (no DB, no network)."""

import pytest

from app.core.exceptions import (
    EmptyFileError,
    FileTooLargeError,
    TextExtractionError,
    UnsupportedFileTypeError,
)
from app.services.chunking import chunk_text
from app.services.extraction import extract_text, validate_upload

# --------------------------------------------------------------------------- #
# Chunking
# --------------------------------------------------------------------------- #


def test_chunks_never_exceed_the_requested_size():
    text = "\n\n".join(f"Paragraph {i}. " + "word " * 80 for i in range(20))

    chunks = chunk_text(text, chunk_size=500, overlap=50)

    assert chunks
    assert all(len(chunk) <= 500 for chunk in chunks)


def test_chunks_overlap_so_boundary_facts_stay_retrievable():
    text = " ".join(f"sentence{i} is a distinct fact here." for i in range(200))

    chunks = chunk_text(text, chunk_size=400, overlap=100)

    assert len(chunks) > 1
    # The start of each chunk should appear in the previous chunk's tail.
    for previous, current in zip(chunks, chunks[1:]):
        assert current[:20] in previous, "expected overlap between adjacent chunks"


def test_short_text_yields_a_single_chunk():
    assert chunk_text("Just one short sentence.", 1000, 100) == [
        "Just one short sentence."
    ]


def test_empty_or_whitespace_text_yields_no_chunks():
    assert chunk_text("", 1000, 100) == []
    assert chunk_text("   \n\n\t  ", 1000, 100) == []


def test_zero_overlap_is_allowed():
    chunks = chunk_text("word " * 500, chunk_size=200, overlap=0)
    assert all(len(chunk) <= 200 for chunk in chunks)


def test_word_longer_than_chunk_size_is_hard_sliced():
    """A single unbroken token must still be split to fit the size limit."""
    chunks = chunk_text("a" * 350, chunk_size=100, overlap=10)

    assert len(chunks) >= 4
    assert all(len(chunk) <= 100 for chunk in chunks)
    # Overlap duplicates characters by design, so total length only has to
    # cover the original rather than equal it exactly.
    assert sum(len(chunk) for chunk in chunks) >= 350
    # Only the original character survives; newlines come from unit separators.
    assert set("".join(chunks).replace("\n", "")) == {"a"}


def test_no_chunk_is_blank():
    text = "Alpha.\n\n\n\nBeta.\n\n\n\n\n\nGamma."
    assert all(chunk.strip() for chunk in chunk_text(text, 20, 5))


@pytest.mark.parametrize("size,overlap", [(0, 0), (-1, 0), (100, 100), (100, 150), (100, -1)])
def test_invalid_chunk_parameters_are_rejected(size, overlap):
    with pytest.raises(ValueError):
        chunk_text("some text", size, overlap)


# --------------------------------------------------------------------------- #
# Validation
# --------------------------------------------------------------------------- #


@pytest.mark.parametrize(
    "filename", ["notes.exe", "archive.zip", "image.png", "legacy.doc", "noextension"]
)
def test_unsupported_extensions_are_rejected(filename):
    with pytest.raises(UnsupportedFileTypeError):
        validate_upload(filename, "application/octet-stream", b"some bytes")


def test_empty_payload_is_rejected():
    with pytest.raises(EmptyFileError):
        validate_upload("notes.txt", "text/plain", b"")


def test_oversized_payload_is_rejected():
    from app.core.config import settings

    with pytest.raises(FileTooLargeError):
        validate_upload("notes.txt", "text/plain", b"x" * (settings.max_upload_bytes + 1))


def test_pdf_without_magic_bytes_is_rejected():
    with pytest.raises(UnsupportedFileTypeError):
        validate_upload("doc.pdf", "application/pdf", b"this is plain text")


def test_docx_without_zip_magic_is_rejected():
    with pytest.raises(UnsupportedFileTypeError):
        validate_upload("doc.docx", None, b"not a zip archive")


def test_valid_extensions_pass_and_normalise_case():
    assert validate_upload("NOTES.TXT", "text/plain", b"hello") == "txt"
    assert validate_upload("doc.PDF", "application/pdf", b"%PDF-1.7 body") == "pdf"


def test_octet_stream_is_tolerated_when_magic_bytes_match():
    """Browsers often send a generic type; magic bytes are the real check."""
    assert validate_upload("doc.pdf", "application/octet-stream", b"%PDF-1.4 x") == "pdf"


# --------------------------------------------------------------------------- #
# Extraction
# --------------------------------------------------------------------------- #


def test_txt_extraction_normalises_whitespace_and_newlines():
    result = extract_text(b"Line one\r\n\r\n\r\n\r\nLine    two", "txt")

    assert result == "Line one\n\nLine two"


def test_latin1_text_is_decoded_rather_than_failing():
    assert "caf" in extract_text("café résumé".encode("latin-1"), "txt")


def test_null_bytes_are_stripped():
    assert "\x00" not in extract_text(b"clean\x00text", "txt")


def test_text_with_no_readable_content_raises():
    with pytest.raises(TextExtractionError):
        extract_text(b"   \n\n\t ", "txt")


def test_docx_extraction_includes_table_cells():
    """Key facts often live in tables, not paragraphs."""
    import io

    from docx import Document as DocxDocument

    document = DocxDocument()
    document.add_paragraph("Contract Summary")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Renewal date"
    table.rows[0].cells[1].text = "1 January 2025"
    buffer = io.BytesIO()
    document.save(buffer)

    result = extract_text(buffer.getvalue(), "docx")

    assert "Contract Summary" in result
    assert "Renewal date" in result
    assert "1 January 2025" in result


def test_corrupt_pdf_raises_extraction_error():
    with pytest.raises(TextExtractionError):
        extract_text(b"%PDF-1.4\ngarbage that is not a real pdf body", "pdf")
