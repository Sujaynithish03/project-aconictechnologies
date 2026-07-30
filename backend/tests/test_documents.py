"""Upload validation, document listing, and tenant isolation."""

import io
import uuid

import pytest
from docx import Document as DocxDocument

from app.api.deps import get_llm_factory
from app.main import app

TXT_BODY = (
    "Acme Corporation Refund Policy\n\n"
    "Customers may request a full refund within 30 days of purchase. "
    "Refunds are processed within 5 business days.\n\n"
    "The policy was last updated on 14 March 2024."
).encode()


@pytest.fixture(autouse=True)
def _use_stub_llm(stub_llm):
    """Route uploads through the stub so no API key or network is needed."""
    app.dependency_overrides[get_llm_factory] = lambda: (lambda: stub_llm)
    yield
    app.dependency_overrides.clear()


def _upload(client, headers, filename, data, content_type):
    return client.post(
        "/upload",
        headers=headers,
        files={"file": (filename, io.BytesIO(data), content_type)},
    )


def _process(client, headers, document_id):
    """Run the embedding phase — the second half of the two-phase upload."""
    return client.post(f"/documents/{document_id}/process", headers=headers)


def _make_docx() -> bytes:
    document = DocxDocument()
    document.add_paragraph("Quarterly Report")
    document.add_paragraph("Revenue grew 18% in Q3 2024.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Deadline"
    table.rows[0].cells[1].text = "31 December 2024"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _make_pdf() -> bytes:
    """Minimal single-page PDF with extractable text, built by hand."""
    from pypdf import PdfWriter

    # pypdf can't author text, so assemble a tiny valid PDF directly.
    content = b"BT /F1 24 Tf 72 700 Td (Invoice total is 4200 USD) Tj ET"
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length " + str(len(content)).encode() + b" >>\nstream\n" + content + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for index, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{index} 0 obj\n".encode() + body + b"\nendobj\n"
    xref_at = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for offset in offsets:
        out += f"{offset:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_at}\n".encode()
        + b"%%EOF\n"
    )
    assert PdfWriter is not None  # sanity: pypdf importable
    return bytes(out)


# --------------------------------------------------------------------------- #
# Validation
# --------------------------------------------------------------------------- #


def test_upload_rejects_unsupported_extension(client, auth_headers):
    response = _upload(client, auth_headers, "malware.exe", b"MZ\x90\x00", "application/octet-stream")

    assert response.status_code == 415
    assert "PDF, DOCX, or TXT" in response.json()["detail"]


def test_upload_rejects_renamed_file_by_magic_bytes(client, auth_headers):
    """A .pdf extension is not enough — the content must actually be a PDF."""
    response = _upload(client, auth_headers, "fake.pdf", b"MZ\x90\x00 not a pdf", "application/pdf")

    assert response.status_code == 415
    assert "valid PDF" in response.json()["detail"]


def test_upload_rejects_empty_file(client, auth_headers):
    response = _upload(client, auth_headers, "empty.txt", b"", "text/plain")
    assert response.status_code == 400


def test_upload_rejects_oversized_file(client, auth_headers):
    from app.core.config import settings

    oversized = b"a" * (settings.max_upload_bytes + 1)
    response = _upload(client, auth_headers, "big.txt", oversized, "text/plain")

    assert response.status_code == 413


def test_upload_rejects_mismatched_content_type(client, auth_headers):
    response = _upload(client, auth_headers, "doc.txt", b"hello", "image/png")
    assert response.status_code == 415


def test_validation_errors_take_precedence_over_missing_llm_key(client, auth_headers):
    """A bad file must report 415, not 503, even with no API key configured.

    Regression: injecting the provider as a route dependency made it resolve
    before the route body, masking every validation error with 503.
    """
    app.dependency_overrides.clear()

    response = _upload(client, auth_headers, "malware.exe", b"MZ\x90\x00", "application/octet-stream")

    assert response.status_code == 415
    assert "PDF, DOCX, or TXT" in response.json()["detail"]


def test_upload_succeeds_without_an_llm_key(client, auth_headers):
    """Upload only extracts text, so it must not depend on the AI service."""
    app.dependency_overrides.clear()

    response = _upload(client, auth_headers, "policy.txt", TXT_BODY, "text/plain")

    assert response.status_code == 202
    assert response.json()["status"] == "pending"
    assert response.json()["char_count"] > 0


def test_process_returns_503_when_llm_is_not_configured(client, auth_headers):
    """The embedding phase is the only part that needs the AI service."""
    document_id = _upload(
        client, auth_headers, "policy.txt", TXT_BODY, "text/plain"
    ).json()["id"]
    app.dependency_overrides.clear()

    response = client.post(f"/documents/{document_id}/process", headers=auth_headers)

    assert response.status_code == 503


def test_upload_requires_authentication(client):
    response = client.post(
        "/upload", files={"file": ("a.txt", io.BytesIO(b"hello"), "text/plain")}
    )
    assert response.status_code == 401


# --------------------------------------------------------------------------- #
# Happy paths — all three supported formats
# --------------------------------------------------------------------------- #


@pytest.mark.parametrize(
    "filename,data_factory,content_type",
    [
        ("policy.txt", lambda: TXT_BODY, "text/plain"),
        ("report.docx", _make_docx,
         "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        ("invoice.pdf", _make_pdf, "application/pdf"),
    ],
)
def test_each_supported_format_uploads_then_indexes(
    client, auth_headers, filename, data_factory, content_type
):
    upload = _upload(client, auth_headers, filename, data_factory(), content_type)

    assert upload.status_code == 202, upload.text
    document = upload.json()
    assert document["filename"] == filename
    # Phase 1 only extracts text — no chunks yet.
    assert document["status"] == "pending"
    assert document["char_count"] > 0
    assert document["chunk_count"] == 0

    processed = _process(client, auth_headers, document["id"])

    assert processed.status_code == 200, processed.text
    body = processed.json()
    assert body["status"] == "ready", body.get("error_message")
    assert body["chunk_count"] > 0


def test_upload_returns_pending_so_status_is_observable(client, auth_headers):
    """The dashboard needs a real pre-indexed state to display."""
    response = _upload(client, auth_headers, "policy.txt", TXT_BODY, "text/plain")

    assert response.status_code == 202
    assert response.json()["status"] == "pending"


def test_upload_rejects_a_pdf_with_no_extractable_text(client, auth_headers):
    """An unreadable PDF is refused up front rather than stored broken."""
    response = _upload(
        client, auth_headers, "scan.pdf", b"%PDF-1.4\n%%EOF\n", "application/pdf"
    )

    assert response.status_code == 400
    assert "pdf" in response.json()["detail"].lower()
    # Extraction runs before the row is created, so nothing unusable is left
    # behind in the user's library.
    assert client.get("/documents", headers=auth_headers).json()["total"] == 0


def test_processing_is_idempotent_and_does_not_duplicate_chunks(client, auth_headers):
    document_id = _upload(
        client, auth_headers, "policy.txt", TXT_BODY, "text/plain"
    ).json()["id"]

    first = _process(client, auth_headers, document_id).json()
    second = _process(client, auth_headers, document_id).json()

    assert first["status"] == second["status"] == "ready"
    assert first["chunk_count"] == second["chunk_count"]


def test_processing_a_failed_document_retries_it(client, auth_headers, stub_llm):
    """A transient embedding failure must be recoverable without re-uploading."""
    document_id = _upload(
        client, auth_headers, "policy.txt", TXT_BODY, "text/plain"
    ).json()["id"]

    def explode(_texts):
        raise RuntimeError("provider is down")

    original = stub_llm.embed_documents
    stub_llm.embed_documents = explode  # type: ignore[method-assign]
    failed = _process(client, auth_headers, document_id).json()
    assert failed["status"] == "failed"
    assert failed["error_message"]

    # The extracted text was kept, so a retry needs no new upload.
    stub_llm.embed_documents = original  # type: ignore[method-assign]
    recovered = _process(client, auth_headers, document_id).json()
    assert recovered["status"] == "ready"
    assert recovered["chunk_count"] > 0


def test_process_requires_authentication_and_ownership(client, auth_headers):
    document_id = _upload(
        client, auth_headers, "policy.txt", TXT_BODY, "text/plain"
    ).json()["id"]

    assert client.post(f"/documents/{document_id}/process").status_code == 401

    other = client.post(
        "/signup",
        json={"email": f"other-{uuid.uuid4().hex[:8]}@example.com", "password": "Password123"},
    ).json()["access_token"]
    response = client.post(
        f"/documents/{document_id}/process",
        headers={"Authorization": f"Bearer {other}"},
    )
    assert response.status_code == 404


def test_documents_lists_only_own_documents_newest_first(client, auth_headers):
    for name in ["first.txt", "second.txt"]:
        _upload(client, auth_headers, name, TXT_BODY, "text/plain")

    response = client.get("/documents", headers=auth_headers)

    assert response.status_code == 200
    body = response.json()
    assert body["total"] == 2
    assert [d["filename"] for d in body["documents"]] == ["second.txt", "first.txt"]


def test_delete_document_removes_it(client, auth_headers):
    document_id = _upload(client, auth_headers, "policy.txt", TXT_BODY, "text/plain").json()["id"]

    assert client.delete(f"/documents/{document_id}", headers=auth_headers).status_code == 200
    assert client.get(f"/documents/{document_id}", headers=auth_headers).status_code == 404


def test_unknown_document_id_returns_404(client, auth_headers):
    response = client.get(f"/documents/{uuid.uuid4()}", headers=auth_headers)
    assert response.status_code == 404


def test_malformed_uuid_returns_422(client, auth_headers):
    response = client.get("/documents/not-a-uuid", headers=auth_headers)
    assert response.status_code == 422


# --------------------------------------------------------------------------- #
# Tenant isolation — the security property that matters most here
# --------------------------------------------------------------------------- #


def test_users_cannot_read_or_delete_each_others_documents(client):
    def register():
        email = f"user-{uuid.uuid4().hex[:10]}@example.com"
        token = client.post(
            "/signup", json={"email": email, "password": "Password123"}
        ).json()["access_token"]
        return {"Authorization": f"Bearer {token}"}

    alice, bob = register(), register()
    document_id = _upload(client, alice, "secret.txt", TXT_BODY, "text/plain").json()["id"]

    assert client.get(f"/documents/{document_id}", headers=bob).status_code == 404
    assert client.delete(f"/documents/{document_id}", headers=bob).status_code == 404
    assert client.get("/documents", headers=bob).json()["total"] == 0

    # Alice's document is untouched.
    assert client.get(f"/documents/{document_id}", headers=alice).status_code == 200
